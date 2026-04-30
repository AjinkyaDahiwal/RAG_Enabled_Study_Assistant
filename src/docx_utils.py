from docx import Document

def extract_docx_content(docx_path):
    """
    Extract text and preserve styles from DOCX.
    Returns a list of paragraphs with style info.
    """
    doc = Document(docx_path)
    content = []
    for para in doc.paragraphs:
        style = para.style.name
        text = para.text
        if text.strip():
            content.append({'style': style, 'text': text})

     # Extract tables
    for tbl_num, table in enumerate(doc.tables):
        table_content = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_content.append(row_data)
        content.append({'type': 'table', 'table_num': tbl_num + 1, 'rows': table_content})
    return content
